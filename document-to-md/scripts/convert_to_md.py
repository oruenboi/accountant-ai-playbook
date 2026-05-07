#!/usr/bin/env python3
"""
convert_to_md.py

Convert common document formats to Markdown with best-effort preservation of
structure. Prefers Pandoc when available, with Python fallbacks so the agent can
still extract usable text when Pandoc is absent.
"""
import argparse
import subprocess
import sys
import shutil
import re
from pathlib import Path
from typing import Optional, Sequence, Tuple, Dict
from datetime import datetime
from collections import defaultdict

SUPPORTED_EXTS = {
    "docx",
    "doc",
    "odt",
    "rtf",
    "html",
    "htm",
    "md",
    "txt",
    "pdf",
    "epub",
    "pptx",
    "xlsx",
    "csv",
    "tsv",
    "jpg",
    "jpeg",
    "png",
    "tif",
    "tiff",
    "gif",
    "bmp",
}


def log(message: str) -> None:
    sys.stderr.write(message + "\n")


def run_cmd(cmd: Sequence[str]) -> subprocess.CompletedProcess:
    return subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def pandoc_exists() -> bool:
    return shutil.which("pandoc") is not None


def convert_with_pandoc(src: Path, dest: Path, media_dir: Optional[Path]) -> bool:
    cmd = ["pandoc", str(src), "-t", "gfm", "-o", str(dest), "--wrap=none"]
    if media_dir:
        cmd += ["--extract-media", str(media_dir)]
    result = run_cmd(cmd)
    if result.returncode != 0:
        log(f"pandoc failed on {src.name}: {result.stderr.strip()}")
        return False
    return True


def markdown_from_pdf(src: Path, allow_ocr: bool) -> Tuple[Optional[str], bool]:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        log("pdfplumber not installed; skipping PDF text extraction")
        return None, False

    texts = []
    with pdfplumber.open(str(src)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            texts.append(page_text)
    joined = "\n\n".join(texts).strip()
    if joined:
        return joined, False

    if not allow_ocr:
        return None, False

    try:
        from pdf2image import convert_from_path  # type: ignore
        import pytesseract  # type: ignore
    except ImportError:
        log("pdf pages had no text; install pdf2image and pytesseract for OCR")
        return None, False

    ocr_chunks = []
    for img in convert_from_path(str(src)):
        ocr_chunks.append(pytesseract.image_to_string(img))
    return "\n\n".join(ocr_chunks).strip(), True


def markdown_from_docx(src: Path) -> Optional[str]:
    try:
        import docx  # type: ignore
    except ImportError:
        log("python-docx not installed; cannot fallback for docx")
        return None

    document = docx.Document(str(src))
    lines = []
    for para in document.paragraphs:
        lines.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n\n".join(line for line in lines if line.strip())


def markdown_from_html(src: Path) -> Optional[str]:
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        log("beautifulsoup4 not installed; cannot fallback for html")
        return None

    soup = BeautifulSoup(src.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for bad in soup(["script", "style"]):
        bad.decompose()
    text = soup.get_text("\n")
    cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    return cleaned


def markdown_from_image(src: Path) -> Optional[str]:
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore
    except ImportError:
        log("pytesseract and pillow required for image OCR")
        return None
    image = Image.open(str(src))
    return pytesseract.image_to_string(image)


def write_markdown(content: str, dest: Path) -> None:
    ensure_parent(dest)
    dest.write_text(content, encoding="utf-8")


def markdown_from_table(src: Path, ext: str) -> Optional[str]:
    try:
        import pandas as pd  # type: ignore
    except ImportError:
        log("pandas not installed; cannot handle tabular file fallback")
        return None

    try:
        if ext == "xlsx":
            df = pd.read_excel(src)
        elif ext == "tsv":
            df = pd.read_csv(src, sep="\t")
        else:  # csv or other
            df = pd.read_csv(src)
    except Exception as exc:  # pragma: no cover - defensive
        log(f"failed reading table {src.name}: {exc}")
        return None

    return df.to_markdown(index=False)


def bold_currency(text: str) -> str:
    return re.sub(r"(\$[0-9][\d,]*(?:\.\d+)?)", r"**\1**", text)


def normalize_definitions(line: str) -> Optional[str]:
    # Pattern: "Term" means ...
    m = re.match(r'^\s*[\"“](.+?)[\"”]\s+means\s+(.*)$', line, flags=re.IGNORECASE)
    if m:
        term, definition = m.group(1), m.group(2)
        return f'- “{term}”: {definition}'
    return None


def legal_normalize(content: str) -> str:
    lines = []
    for raw in content.splitlines():
        line = raw.replace("\u00a0", " ")  # NBSP to space
        # Standardize (a)/(b) bullets to "- (a) ..."
        line = re.sub(r"^\s*\(([a-z])\)\s+", r"- (\1) ", line)
        # Try to normalize definitions
        norm_def = normalize_definitions(line)
        if norm_def is not None:
            line = norm_def
        # Highlight currency
        line = bold_currency(line)
        lines.append(line)

    # Flatten wrapped bullet lines
    bullet_re = re.compile(r"^\s*(?:-|\*|\+|\d+\.)\s|^\s*\([a-z]\)")
    flattened: list[str] = []
    for line in lines:
        if not flattened:
            flattened.append(line)
            continue
        if line.strip() and not bullet_re.match(line) and bullet_re.match(flattened[-1]):
            flattened[-1] = flattened[-1].rstrip() + " " + line.strip()
        else:
            flattened.append(line)

    # Collapse multiple blank lines
    cleaned: list[str] = []
    blank = 0
    for line in flattened:
        if line.strip() == "":
            blank += 1
            if blank > 1:
                continue
        else:
            blank = 0
        cleaned.append(line.rstrip())

    return "\n".join(cleaned)


def extract_headings(content: str) -> list[Tuple[int, str]]:
    headings = []
    for line in content.splitlines():
        m = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            headings.append((level, text))
    return headings


def add_metadata_block(
    content: str,
    *,
    source: Path,
    ext: str,
    used_pandoc: bool,
    used_ocr: bool,
    media_dir: Optional[Path],
    heading_count: int,
    line_count: int,
    add_meta: bool,
    meta_extras: Dict[str, str],
) -> str:
    if not add_meta:
        return content
    # Avoid duplicating if the content already begins with frontmatter.
    stripped = content.lstrip()
    if stripped.startswith("---"):
        return content

    meta_lines = [
        "---",
        f"source: {source}",
        f"original_extension: .{ext}",
        f"converted_at: {datetime.now().astimezone().isoformat()}",
        "converter: document-to-md",
        f"used_pandoc: {str(used_pandoc).lower()}",
        f"ocr_used: {str(used_ocr).lower()}",
        f"heading_count: {heading_count}",
        f"line_count: {line_count}",
    ]
    for key, val in meta_extras.items():
        if val:
            meta_lines.append(f"{key}: {val}")
    if media_dir:
        meta_lines.append(f"media_dir: {media_dir}")
    meta_lines.append("---")
    meta_lines.append("")  # blank line after metadata
    return "\n".join(meta_lines) + content


def insert_toc(
    content: str,
    headings: list[Tuple[int, str]],
    *,
    force: bool,
    disable: bool,
    threshold: int,
) -> str:
    if disable:
        return content
    line_count = content.count("\n") + 1
    if not force:
        if line_count < threshold or len(headings) < 3:
            return content

    slug_counts: defaultdict[str, int] = defaultdict(int)
    toc_lines = []
    for level, text in headings:
        slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
        slug_counts[slug] += 1
        if slug_counts[slug] > 1:
            slug = f"{slug}-{slug_counts[slug]-1}"
        indent = "  " * (level - 1)
        toc_lines.append(f"{indent}- [{text}](#{slug})")

    toc_block = ["## Table of Contents"] + toc_lines + [""]

    lines = content.splitlines()
    insert_at = 0
    if lines and lines[0].strip() == "---":
        for idx in range(1, len(lines)):
            if lines[idx].strip() == "---":
                insert_at = idx + 1
                break
    new_lines = lines[:insert_at] + toc_block + lines[insert_at:]
    return "\n".join(new_lines)


def convert_one(
    src: Path,
    out_dir: Optional[Path],
    media_dir: Optional[Path],
    allow_ocr: bool,
    add_meta: bool,
    force_toc: bool,
    disable_toc: bool,
    toc_threshold: int,
    legal_mode: bool,
    meta_extras: Dict[str, str],
) -> bool:
    if not src.exists():
        log(f"missing file: {src}")
        return False

    ext = src.suffix.lower().lstrip(".")
    if ext not in SUPPORTED_EXTS:
        log(f"unsupported extension: .{ext}")
        return False

    dest_dir = out_dir if out_dir else src.parent
    dest = dest_dir / f"{src.stem}.md"

    used_pandoc = False
    used_ocr = False
    content: Optional[str] = None

    # Pandoc is preferred for structured formats.
    if pandoc_exists() and ext not in {"jpg", "jpeg", "png", "tif", "tiff", "gif", "bmp"}:
        if convert_with_pandoc(src, dest, media_dir):
            used_pandoc = True
            log(f"pandoc conversion succeeded: {src} -> {dest}")
            try:
                content = dest.read_text(encoding="utf-8", errors="ignore")
            except Exception as exc:  # pragma: no cover
                log(f"failed to read pandoc output for post-processing: {exc}")
                return False

    if content is None:
        if ext in {"pdf"}:
            content, used_ocr = markdown_from_pdf(src, allow_ocr)
        elif ext in {"docx", "doc"}:
            content = markdown_from_docx(src)
        elif ext in {"html", "htm"}:
            content = markdown_from_html(src)
        elif ext in {"xlsx", "csv", "tsv"}:
            content = markdown_from_table(src, ext)
            if content is None and ext in {"csv", "tsv"}:
                content = src.read_text(encoding="utf-8", errors="ignore")
        elif ext in {"txt", "md"}:
            content = src.read_text(encoding="utf-8", errors="ignore")
        elif ext in {"jpg", "jpeg", "png", "tif", "tiff", "gif", "bmp"}:
            content = markdown_from_image(src)
        else:
            log(f"no fallback implemented for .{ext}")

    if not content:
        log(f"failed to extract text from {src}")
        return False

    if legal_mode:
        content = legal_normalize(content)

    headings = extract_headings(content)
    line_count = content.count("\n") + 1
    content = add_metadata_block(
        content,
        source=src,
        ext=ext,
        used_pandoc=used_pandoc,
        used_ocr=used_ocr,
        media_dir=media_dir,
        heading_count=len(headings),
        line_count=line_count,
        add_meta=add_meta,
        meta_extras=meta_extras,
    )

    # Recompute headings after metadata (metadata shouldn't add headings).
    content = insert_toc(
        content,
        headings,
        force=force_toc,
        disable=disable_toc,
        threshold=toc_threshold,
    )

    write_markdown(content, dest)
    log(f"wrote markdown: {dest}")
    return True


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert documents to Markdown")
    parser.add_argument("inputs", nargs="+", help="Files to convert")
    parser.add_argument("--out-dir", type=Path, help="Directory for .md outputs")
    parser.add_argument(
        "--media-dir",
        type=Path,
        help="Where to extract embedded media when supported",
    )
    parser.add_argument(
        "--ocr",
        action="store_true",
        help="Allow OCR when PDFs are image-based",
    )
    parser.add_argument(
        "--no-meta",
        action="store_true",
        help="Skip adding YAML metadata header",
    )
    parser.add_argument(
        "--toc",
        action="store_true",
        help="Force a Table of Contents insertion",
    )
    parser.add_argument(
        "--no-toc",
        action="store_true",
        help="Disable automatic Table of Contents insertion",
    )
    parser.add_argument(
        "--toc-threshold",
        type=int,
        default=120,
        help="Minimum line count to auto-insert TOC when headings exist (default: 120)",
    )
    parser.add_argument(
        "--legal-normalize",
        action="store_true",
        help="Apply legal-text friendly normalizations (bullet cleanup, unwrap, currency bold, definition bullets)",
    )
    parser.add_argument("--meta-date-in-force", help="Add date_in_force to metadata header (e.g., 2025-06-09)")
    parser.add_argument("--meta-made-on", help="Add made_on to metadata header")
    parser.add_argument("--meta-citation", help="Add citation to metadata header")
    parser.add_argument("--meta-source-url", help="Add source_url to metadata header")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    meta_extras = {
        "date_in_force": args.meta_date_in_force or "",
        "made_on": args.meta_made_on or "",
        "citation": args.meta_citation or "",
        "source_url": args.meta_source_url or "",
    }
    success = True
    for raw in args.inputs:
        src = Path(raw).expanduser().resolve()
        ok = convert_one(
            src,
            args.out_dir,
            args.media_dir,
            args.ocr,
            add_meta=not args.no_meta,
            force_toc=args.toc,
            disable_toc=args.no_toc,
            toc_threshold=args.toc_threshold,
            legal_mode=args.legal_normalize,
            meta_extras=meta_extras,
        )
        success = success and ok
    return 0 if success else 1


if __name__ == "__main__":
    raise SystemExit(main())
