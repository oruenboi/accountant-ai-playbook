#!/usr/bin/env python3
"""Generate tiny fixture documents for smoke tests."""
from pathlib import Path
from docx import Document
from fpdf import FPDF

FIXTURES_DIR = Path(__file__).resolve().parent.parent / "assets" / "fixtures"
FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

# DOCX
doc = Document()
doc.add_heading("Sample Contract Excerpt", level=1)
doc.add_paragraph("1. Grant of access. The Service Provider shall grant access to the system.")
doc.add_paragraph("2. Fees. Licensee shall pay $10,000 per annum.")
doc.save(FIXTURES_DIR / "sample.docx")

# PDF
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.multi_cell(
    0,
    10,
    "Sample PDF\n\nPart 1 - Preliminary\n(a) Definitions. \"Agreement\" means this contract.\n"
    "(b) Payment. Fee is $250.\n\nPart 2 - Operations\n1. Access to electronic transaction system.",
)
pdf.output(str(FIXTURES_DIR / "sample.pdf"))

# HTML
html = """<!doctype html>
<html><head><title>Sample HTML</title></head>
<body>
<h1>Sample Policy</h1>
<h2>Section 3. Access</h2>
<p>Customer may access the system subject to payment of <strong>$500</strong>.</p>
</body></html>
"""
(FIXTURES_DIR / "sample.html").write_text(html, encoding="utf-8")

# CSV
(FIXTURES_DIR / "sample.csv").write_text("Name,Fee\nFiling,100\nProcessing,250\n", encoding="utf-8")

# TXT
(FIXTURES_DIR / "sample.txt").write_text(
    "Plain text sample for fallback. Bullet: (a) first point; (b) second point.",
    encoding="utf-8",
)

print(f"Fixtures written to {FIXTURES_DIR}")
